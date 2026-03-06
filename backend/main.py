from __future__ import annotations

from io import BytesIO
import re
from typing import Optional

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from claude_service import ClaudeResponseError, ClaudeService
from models import CardRequest, CardResponse, ExportRequest, QuoteRequest, QuotesResponse, SummaryResponse, UrlRequest
from scraper import ScrapeError, fetch_article

app = FastAPI(title="CutCards API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


_claude: Optional[ClaudeService] = None


def get_claude() -> ClaudeService:
    global _claude
    if _claude is None:
        _claude = ClaudeService()
    return _claude


@app.get("/health")
def health() -> dict:
    return {"ok": True}


@app.post("/api/summary", response_model=SummaryResponse)
def summarize(payload: UrlRequest) -> SummaryResponse:
    try:
        article = fetch_article(str(payload.url))
        summary = get_claude().summarize_article(article, payload.debate_topic.strip())
        return SummaryResponse(article=article, summary=summary)
    except ScrapeError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except ClaudeResponseError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Unexpected error: {exc}") from exc


@app.post("/api/quotes", response_model=QuotesResponse)
def quotes(payload: QuoteRequest) -> QuotesResponse:
    try:
        items = get_claude().extract_relevant_quotes(
            payload.article,
            payload.debate_topic.strip(),
            payload.custom_instructions.strip(),
        )
        return QuotesResponse(quotes=items)
    except ClaudeResponseError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Unexpected error: {exc}") from exc


@app.post("/api/card", response_model=CardResponse)
def build_card(payload: CardRequest) -> CardResponse:
    try:
        return get_claude().build_card(
            payload.article,
            payload.debate_topic.strip(),
            payload.selected_quotes,
            payload.custom_instructions.strip(),
        )
    except ClaudeResponseError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Unexpected error: {exc}") from exc


def _parse_cite(cite: str, source_url: Optional[str]) -> tuple[str, str]:
    cleaned = " ".join(cite.split())
    parts = [part.strip() for part in cleaned.split(",") if part.strip()]
    lead = parts[0] if parts else "Unknown"
    lead_last_name = lead.split()[-1]
    year_match = re.search(r"(19|20)\d{2}", cleaned)
    year = year_match.group(0) if year_match else ""
    lead_chunk = f"{lead_last_name} {year}".strip()

    details = cleaned
    if source_url and source_url not in details:
        details = f"{details} {source_url}".strip()
    return lead_chunk, details


def _normalize_for_match(text: str) -> str:
    return (
        text.replace("“", '"')
        .replace("”", '"')
        .replace("’", "'")
        .replace("‘", "'")
        .replace("–", "-")
        .replace("—", "-")
    )


def _find_quote_span(paragraph_text: str, quote_text: str) -> tuple[int, int] | None:
    normalized_paragraph = _normalize_for_match(paragraph_text)
    normalized_quote = _normalize_for_match(quote_text).strip().strip('"').strip("'")
    tokens = normalized_quote.split()
    if not tokens:
        return None

    pattern = re.compile(r"\s+".join(re.escape(token) for token in tokens), flags=re.IGNORECASE)
    match = pattern.search(normalized_paragraph)
    if not match:
        return None
    return match.start(), match.end()


def _style_run(run, *, bold: bool, underline: bool, size: int, color: RGBColor, highlight=None) -> None:
    run.font.name = "Times New Roman"
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight is not None:
        run.font.highlight_color = highlight


def _append_context_and_quote(paragraph, paragraph_text: str, quote_text: str, primary: bool, format_variant: str) -> None:
    span = _find_quote_span(paragraph_text, quote_text)
    size_map = {"classic": (14, 15), "spotlight": (13, 16), "brief": (13, 14)}
    context_size, quote_size = size_map.get(format_variant, (14, 15))
    if primary:
        context_size += 1
        quote_size += 1

    context_bold = format_variant != "brief"
    context_color = RGBColor(0, 0, 0)
    quote_highlight = WD_COLOR_INDEX.YELLOW if format_variant == "spotlight" else WD_COLOR_INDEX.BRIGHT_GREEN

    if not span:
        context_run = paragraph.add_run(paragraph_text.strip())
        _style_run(
            context_run,
            bold=context_bold,
            underline=True,
            size=context_size,
            color=context_color,
        )
        paragraph.add_run(" ")
        open_quote = paragraph.add_run('"')
        _style_run(open_quote, bold=True, underline=True, size=quote_size, color=context_color)
        quote_run = paragraph.add_run(quote_text.strip().strip('"').strip("'"))
        _style_run(
            quote_run,
            bold=True,
            underline=True,
            size=quote_size,
            color=context_color,
            highlight=quote_highlight,
        )
        close_quote = paragraph.add_run('"')
        _style_run(close_quote, bold=True, underline=True, size=quote_size, color=context_color)
        return

    start, end = span
    before = paragraph_text[:start]
    quote = paragraph_text[start:end]
    after = paragraph_text[end:]

    if before.endswith(("“", '"', "‘", "'")) and after.startswith(("”", '"', "’", "'")):
        before = before[:-1]
        after = after[1:]
        quote = quote.strip().strip('"').strip("'")

    if before:
        before_run = paragraph.add_run(before)
        _style_run(
            before_run,
            bold=context_bold,
            underline=True,
            size=context_size,
            color=context_color,
        )

    open_quote = paragraph.add_run('"')
    _style_run(open_quote, bold=True, underline=True, size=quote_size, color=context_color)

    quote_run = paragraph.add_run(quote.strip().strip('"').strip("'"))
    _style_run(
        quote_run,
        bold=True,
        underline=True,
        size=quote_size,
        color=context_color,
        highlight=quote_highlight,
    )

    close_quote = paragraph.add_run('"')
    _style_run(close_quote, bold=True, underline=True, size=quote_size, color=context_color)

    if after:
        after_run = paragraph.add_run(after)
        _style_run(
            after_run,
            bold=context_bold,
            underline=True,
            size=context_size,
            color=context_color,
        )


@app.post("/api/export")
def export_card(payload: ExportRequest) -> StreamingResponse:
    if not payload.quote_blocks:
        raise HTTPException(status_code=400, detail="At least one quote is required for export")

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    tag_paragraph = doc.add_paragraph()
    tag_paragraph.paragraph_format.space_before = Pt(0)
    tag_paragraph.paragraph_format.space_after = Pt(0)
    tag_run = tag_paragraph.add_run(payload.tag.strip())
    tag_run.bold = True
    tag_run.underline = True
    tag_run.font.size = Pt(18)
    tag_run.font.color.rgb = RGBColor(0, 0, 0)

    lead, details = _parse_cite(payload.cite, payload.source_url)
    cite_paragraph = doc.add_paragraph()
    cite_paragraph.paragraph_format.space_before = Pt(0)
    cite_paragraph.paragraph_format.space_after = Pt(0)
    lead_run = cite_paragraph.add_run(lead)
    lead_run.bold = True
    lead_run.font.size = Pt(15)
    lead_run.font.color.rgb = RGBColor(0, 0, 0)
    details_run = cite_paragraph.add_run(f" {details}")
    details_run.font.size = Pt(9)
    details_run.font.color.rgb = RGBColor(105, 105, 105)

    primary_index = payload.primary_warrant_quote_index
    if primary_index < 0 or primary_index >= len(payload.quote_blocks):
        primary_index = 0

    for quote_index, block in enumerate(payload.quote_blocks):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        _append_context_and_quote(
            p,
            block.paragraph_text.strip(),
            block.quote_text.strip(),
            primary=quote_index == primary_index,
            format_variant=payload.format_variant,
        )

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    headers = {"Content-Disposition": "attachment; filename=cutcard.docx"}
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
